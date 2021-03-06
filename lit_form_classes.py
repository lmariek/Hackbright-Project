"""FACTORY OF LITITGATION FORM OBJECTS."""

from textblob import TextBlob
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from defenses import all_defenses
import inflect


class ComplaintForm(TextBlob):
    def __init__(self, decoded_text):
        super(ComplaintForm, self).__init__(decoded_text)

        self.word_list = self.words
        self.sentence_list = self.sentences
        self.nouns = self.noun_phrases
        self.plaintiff_fname = self.get_plaintiff_fname()
        self.plaintiff_lname = self.get_plaintiff_lname()
        self.case_no = self.get_case_no()
        self.county = self.get_county()
        self.state = self.get_state()
        self.defendant_fname = self.get_defendant_fname()
        self.defendant_lname = self.get_defendant_lname()
        self.defendant_residence = self.get_defendant_residence()
        self.amount_claimed = self.get_amount_claimed()
        self.claim = self.get_claim()
        self.counsel_fname = self.get_counsel_fname()
        self.counsel_lname = self.get_counsel_lname()
        self.counsel_firm = self. get_counsel_firm()
        self.complaint_date = self.get_complaint_date()
        self.legal_basis = self.get_legal_basis()

    def get_plaintiff_fname(self):
        """Return the plaintiff's first name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'Plaintiff':
                plaintiff_fname = self.word_list[i - 4]

                return plaintiff_fname.capitalize()


    def get_plaintiff_lname(self):
        """Return the plaintiff's last name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'Plaintiff':
                plaintiff_lname = self.word_list[i - 3]

                return plaintiff_lname.capitalize()


    def get_case_no(self):
        """Return the plaintiff's last name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'No':
                case_no = self.word_list[i + 1]

                return case_no


    def get_county(self):
        """Return the county."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'COUNTY':
                if self.word_list[i + 1] == 'OF':
                    county = self.word_list[i + 2]

                    if self.word_list[i + 3].isupper():
                        county = county + ' ' + self.word_list[i + 3]

                    return county.title()


    def get_state(self):
        """Return the state."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'STATE':
                if self.word_list[i + 1] == 'OF':
                    state = self.word_list[i + 2]

                    if self.word_list[i + 3].isupper():
                        state = state + ' ' + self.word_list[i + 3]

                    return state.title()


    def get_defendant_fname(self):
        """Return the defendant's first name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'defendant':

                return self.word_list[i + 1]


    def get_defendant_lname(self):
        """Return the defendant's last name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'defendant':

                return self.word_list[i + 2]


    def get_amount_claimed(self):
        """Return the dollar of damages requested."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'Amount':
                if self.word_list[i + 1] == 'claimed':

                    return self.word_list[i + 2]


    def get_claim(self):
        """Return the type of claim as either the default PI or an error message."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('Personal Injury'):
                return 'Personal Injury'
            else:
                return 'UNKNOWN'


    def get_legal_basis(self):
        """Return the types of legal claims made in the complaint.

        TODO: (2.0) allow multiple legal bases."""

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'claim' and self.nouns[i + 1] == 'relief':
                result = self.nouns[i + 2]

                return result.capitalize()


    def get_defendant_residence(self):
        """Return the defendant's city, county, state of residence."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'resides':
                defendant_residence = self.word_list[i + 2:i + 6]

                return ' '.join(defendant_residence)


    def get_counsel_fname(self):
        """Return the opposing counsel's first name."""

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'llc':
                full_name = self.nouns[i + 1]
                counsel_fname = full_name.split(' ')[0]

                return counsel_fname.capitalize()


    def get_counsel_lname(self):
        """Return the opposing counsel's last name."""

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'llc':
                full_name = self.nouns[i + 1]
                counsel_lname = full_name.split(' ')[2]

                return counsel_lname.capitalize()


    def get_counsel_firm(self):
        """Return the opposing counsel's firm name."""

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'llc':
                firm = self.nouns[i -1]
                firm = firm.title()
                org = self.nouns[i]
                org = org.upper()
                return firm + ' ' + org


    def get_complaint_date(self):
        """Stampe the complaint with the date processed."""

        return datetime.utcnow()


class AnswerForm(object):

    plaintiff_fname = plaintiff_lname = defendant_fname = defendant_lname = None
    case_state = case_county = user_fname = user_lname = user_mailing_address = None
    user_email = user_firm_name = case_no = None

    def __init__(self, complaint, user, defenses):

        self.complaint = complaint
        self.user = user
        self.defenses = defenses
        self.plaintiff_fname = complaint.case.plaintiffs[0].fname
        self.plaintiff_lname = complaint.case.plaintiffs[0].lname
        self.defendant_fname = complaint.case.defendants[0].fname
        self.defendant_lname = complaint.case.defendants[0].lname
        self.case_county = complaint.case.county
        self.case_state = complaint.case.state
        self.user_fname = user.fname
        self.user_lname = user.lname
        self.user_email = user.email
        self.user_mailing_address = user.mailing_address
        self.user_firm_name = user.firm_name
        self.case_no = str(complaint.case.case_no)


    def insert_information(self):
        """Adds custom information into the answer template.

        Returns a modified docx file."""

        # Make a list of all attributes on an Answer class
        attrs = [attr for attr in AnswerForm.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(AnswerForm.__dict__[attr]))]

        # Make a list of all upper case attributes
        cap_attrs = [attr.upper() for attr in AnswerForm.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(AnswerForm.__dict__[attr]))]

        # Make a Document object for Python-docx
        answer = Document('forms/answer_template.docx')
        style = answer.styles['Normal']
        font = style.font
        font.size = Pt(12)

        # Preserve format: replace lowercase tags with values
        for attr_name in attrs:

            for paragraph in answer.paragraphs:

                if attr_name in paragraph.text:
                    paragraph.text = (
                        paragraph.text.replace(attr_name,
                                               getattr(self, attr_name)))

        # Preserve format: replace uppercase tags with values
        for attr_name in cap_attrs:

            for paragraph in answer.paragraphs:

                if attr_name in paragraph.text:
                    attr_data = getattr(self, attr_name.lower())
                    attr_data = attr_data.upper()
                    paragraph.text = (
                        paragraph.text.replace(attr_name, attr_data))

        counter1 = 1
        counter2 = 2
        # TODO: make this its own method on the class
        for defense in self.defenses:

            for p in answer.paragraphs:
                # Begin making paragraphs
                if '***' in p.text:
                    # Grab the legalese from the dictionary - a string of text
                    legalese = all_defenses[defense]

                    # Use inflect library to convert int to spelled digit
                    convert = inflect.engine()
                    spell_digit = convert.number_to_words(counter1)
                    spell_ordinal = convert.ordinal(spell_digit)
                    spell_ordinal = spell_ordinal.upper()

                    # Add the header for each paragraph
                    prior_paragraph = p.insert_paragraph_before()
                    prior_paragraph.add_run("{spell_ordinal} AFFIRMATIVE DEFENSE".format(spell_ordinal=spell_ordinal)).bold = True
                    prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    prior_paragraph = p.insert_paragraph_before()
                    counter1 += 1

                    # Grammar check
                    if spell_ordinal != 'eighth' or 'eleventh':
                        custom_intro = "As a {spell_ordinal}, separate, and affirmative defense".format(spell_ordinal=spell_ordinal.lower())
                    else:
                        custom_intro = "As an {spell_ordinal}, separate, and affirmative defense".format(spell_ordinal=spell_ordinal.lower())

                    # Add the meat of the legal text, with formatting
                    legalese = str(counter2) +'.' + '\t' + '\t' + custom_intro + legalese
                    prior_paragraph = p.insert_paragraph_before(legalese)
                    paragraph_format = prior_paragraph.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.25)
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    counter2 += 1

        # Make a new filename from case no.
        filename = 'answer_{case_no}.docx'.format(case_no=self.case_no)
        # Save the modified document with the new filename
        answer.save('filestorage/{filename}'.format(filename=filename))
        # Return filename to pass to display
        return filename


class InterrogatoriesForm(object):

    plaintiff_fname = plaintiff_lname = defendant_fname = defendant_lname = None
    case_state = case_county = user_fname = user_lname = user_mailing_address = None
    user_email = user_firm_name = case_no = None

    def __init__(self, complaint, user):

        self.complaint = complaint
        self.user = user
        self.plaintiff_fname = complaint.case.plaintiffs[0].fname
        self.plaintiff_lname = complaint.case.plaintiffs[0].lname
        self.defendant_fname = complaint.case.defendants[0].fname
        self.defendant_lname = complaint.case.defendants[0].lname
        self.case_county = complaint.case.county
        self.case_state = complaint.case.state
        self.user_fname = user.fname
        self.user_lname = user.lname
        self.user_email = user.email
        self.user_mailing_address = user.mailing_address
        self.user_firm_name = user.firm_name
        self.case_no = str(complaint.case.case_no)


    def insert_information(self):
        """Adds custom information into the interrogatories template.

        Returns a modified docx file."""

        # Make a list of all attributes on an Interrogatories class
        attrs = [attr for attr in Interrogatories.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(Interrogatories.__dict__[attr]))]

        # Make a list of all upper case attributes
        cap_attrs = [attr.upper() for attr in Interrogatories.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(Interrogatories.__dict__[attr]))]

        # Make a Document object for Python-docx
        interrogatories = Document('forms/interrogatories_from_D_template.docx')
        style = interrogatories.styles['Normal']
        font = style.font
        font.size = Pt(12)

        # Preserve format: replace lowercase tags with values
        for attr_name in attrs:

            for paragraph in answer.paragraphs:

                if attr_name in paragraph.text:
                    paragraph.text = (
                        paragraph.text.replace(attr_name,
                                               getattr(self, attr_name)))

        # Preserve format: replace uppercase tags with values
        for attr_name in cap_attrs:

            for paragraph in answer.paragraphs:

                if attr_name in paragraph.text:
                    attr_data = getattr(self, attr_name.lower())
                    attr_data = attr_data.upper()
                    paragraph.text = (
                        paragraph.text.replace(attr_name, attr_data))

        # Make a new filename from case no.
        filename = 'interrogatories_{case_no}.docx'.format(case_no=self.case_no)
        # Save the modified document with the new filename
        interrogatories.save('filestorage/{filename}'.format(filename=filename))
        # Return filename to pass to display
        return filename


class RequestProDocsForm(object):

    plaintiff_fname = plaintiff_lname = defendant_fname = defendant_lname = None
    case_state = case_county = user_fname = user_lname = user_mailing_address = None
    user_email = user_firm_name = case_no = None

    def __init__(self, complaint, user):

        self.complaint = complaint
        self.user = user
        self.plaintiff_fname = complaint.case.plaintiffs[0].fname
        self.plaintiff_lname = complaint.case.plaintiffs[0].lname
        self.defendant_fname = complaint.case.defendants[0].fname
        self.defendant_lname = complaint.case.defendants[0].lname
        self.case_county = complaint.case.county
        self.case_state = complaint.case.state
        self.user_fname = user.fname
        self.user_lname = user.lname
        self.user_email = user.email
        self.user_mailing_address = user.mailing_address
        self.user_firm_name = user.firm_name
        self.case_no = str(complaint.case.case_no)


    def insert_information(self):
        """Adds custom information into the request for production of documents template.

        Returns a modified docx file."""

        # Make a list of all attributes on a RequestProDocs class
        attrs = [attr for attr in RequestProDocs.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(RequestProDocs.__dict__[attr]))]

        # Make a list of all upper case attributes
        cap_attrs = [attr.upper() for attr in RequestProDocs.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(RequestProDocs.__dict__[attr]))]

        # Make a Document object for Python-docx
        request_pro_docs = Document('forms/request_for_production_of_docs_template.docx')
        style = request_pro_docs.styles['Normal']
        font = style.font
        font.size = Pt(12)

        # Preserve format: replace lowercase tags with values
        for attr_name in attrs:

            for paragraph in request_pro_docs.paragraphs:

                if attr_name in paragraph.text:
                    paragraph.text = (
                        paragraph.text.replace(attr_name,
                                               getattr(self, attr_name)))

        # Preserve format: replace uppercase tags with values
        for attr_name in cap_attrs:

            for paragraph in request_pro_docs.paragraphs:

                if attr_name in paragraph.text:
                    attr_data = getattr(self, attr_name.lower())
                    attr_data = attr_data.upper()
                    paragraph.text = (
                        paragraph.text.replace(attr_name, attr_data))

        # Make a new filename from case no.
        filename = 'request_pro_docs_{case_no}.docx'.format(case_no=self.case_no)
        # Save the modified document with the new filename
        request_pro_docs.save('filestorage/{filename}'.format(filename=filename))
        # Return filename to pass to display
        return filename
